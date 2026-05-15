import io
import os
import time
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm
from dotenv import load_dotenv
from app.chart_gen import insert_visualization

load_dotenv()

OPENROUTER_API_KEY_DEFAULT = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "meta-llama/llama-3.3-70b-instruct:free")

# ===================== ПРИМЕР СТИЛЯ (FEW-SHOT) =====================

STYLE_EXAMPLE = """
Примеры аналитических фрагментов из университетского отчёта:

[Пример 1 — позитивный результат с нюансами]
В текущем учебном году отмечается устойчивая тенденция к росту удовлетворённости
студентов организацией образовательного процесса. Особенно высоко оцениваются
логическая последовательность изучения дисциплин и актуальность получаемых знаний —
это свидетельствует о том, что кафедры последовательно работают над содержанием
программ. Вместе с тем картина не является однородной: заметная часть опрошенных
указывает на нехватку практической составляющей и негибкость расписания. Такое
сочетание сигнализирует о точечных проблемах на фоне общего благополучия — именно
они заслуживают первоочередного внимания. Можно предположить, что адресная работа
по этим направлениям позволит закрепить и усилить положительную динамику.

[Пример 2 — тревожный результат, требующий осмысления]
Результаты опроса фиксируют низкую вовлечённость студентов во внеучебную жизнь
университета: большинство респондентов не участвуют ни в научной, ни в творческой,
ни в волонтёрской деятельности. Это не просто статистика — за ней стоит
сформировавшаяся дистанция между студентом и университетом как сообществом.
Вероятной причиной служит не отсутствие интереса как такового, а недостаточная
информированность и отсутствие первого опыта участия. Показательно, что значительная
часть опрошенных готова включиться при наличии активной информационной поддержки —
это говорит о скрытом потенциале, который пока не задействован. Полученные данные
указывают на необходимость системной работы по вовлечению, а не разовых акций.

[Пример 3 — противоречивый результат]
Распределение ответов на данный вопрос обнаруживает внутреннее противоречие,
заслуживающее отдельного осмысления. С одной стороны, большинство студентов
в целом позитивно оценивают условия обучения; с другой — конкретные аспекты
организации учебного процесса получают заметно более сдержанные оценки.
Подобное расхождение между общей и частной оценкой характерно для ситуаций,
когда негативный опыт ещё не успел сложиться в устойчивое критическое отношение.
Вместе с тем игнорировать эти сигналы не следует: именно из таких частностей
со временем формируется общая неудовлетворённость. Университету стоит
рассмотреть эти зоны как точки превентивного вмешательства.
"""


# ===================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =====================

def _p(doc, text, bold=False, size=12, space_before=0, space_after=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return para


def _make_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    return doc


# ===================== ДАННЫЕ (файл 1) =====================

def generate_data_docx(questions: list) -> bytes:
    """Файл со списками вопросов и статистикой без аналитики."""
    doc = _make_doc()
    _last_section = None

    for q in questions:
        sec = q.get("section")
        sec_name = sec.get("name") if sec else None

        if sec_name and sec_name != _last_section:
            _last_section = sec_name
            if doc.paragraphs:
                doc.add_page_break()
            _p(doc, sec_name, bold=True, size=14, space_after=4)
            if sec.get("description"):
                _p(doc, sec["description"], size=12, space_after=8)

        file_keys = q["file_keys"]
        file_labels = q["file_labels"]
        file_totals = q["file_totals"]
        is_single = len(file_keys) == 1

        _p(doc, f"Вопрос {q['table_num']} – «{q['question_name']}»",
           bold=True, space_before=10, space_after=2)

        for row in q["rows"]:
            if is_single:
                fk = file_keys[0]
                count = row["counts"].get(fk, 0)
                total = file_totals.get(fk, 0)
                pct = f"{count / total * 100:.1f}%" if total > 0 else "—"
                _p(doc, f"  • {row['answer']}: {count} ({pct})", space_after=1)
            else:
                parts = []
                for fk in file_keys:
                    label = file_labels.get(fk, fk)
                    count = row["counts"].get(fk, 0)
                    total = file_totals.get(fk, 0)
                    pct = f"{count / total * 100:.1f}%" if total > 0 else "—"
                    parts.append(f"{label}: {count} ({pct})")
                _p(doc, f"  • {row['answer']}: {'; '.join(parts)}", space_after=1)

        if q.get("show_total", True):
            if is_single:
                fk = file_keys[0]
                _p(doc, f"  Всего: {file_totals.get(fk, 0)}", bold=True, space_after=6)
            else:
                parts = [
                    f"{file_labels.get(fk, fk)}: {file_totals.get(fk, 0)}"
                    for fk in file_keys
                ]
                _p(doc, f"  Всего: {'; '.join(parts)}", bold=True, space_after=6)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ===================== ПРОМПТЫ =====================

def _build_question_prompt(
    question: dict,
    sec_name: str = "",
    sec_description: str = "",
) -> tuple[str, str]:

    # ── SYSTEM ──────────────────────────────────────────────────────────
    system_parts = [
        "Ты — аналитик социологических исследований в Мордовском Государственном Университете им Н.П. Огарёва.",
        "Пиши официальные аналитические фрагменты для отчётов.",
    ]

    if sec_description:
        system_parts += [
            "",
            "## Контекст текущего раздела",
            sec_description,
            "",
            "Используй этот контекст при написании анализа:",
            "— Если здесь сформулировано конкретное требование или акцент — строго следуй ему.",
            "— Если это описание тематики или состава раздела — учитывай при интерпретации.",
            "— Если это пояснение или замечание — прими во внимание как фоновый контекст.",
            "Академический стиль — инструмент, он не должен вступать в противоречие с контекстом раздела.",
        ]

    system_prompt = "\n".join(system_parts)

    # ── USER ─────────────────────────────────────────────────────────────
    lines = [
        "Ниже — примеры желаемого стиля:",
        STYLE_EXAMPLE,
        "",
        "ПРАВИЛА НАПИСАНИЯ:",
        "— Не пересказывай статистику подряд — интерпретируй, что за ней стоит.",
        "— Ищи противоречия, неожиданности, скрытые смыслы в данных.",
        "— Дай результату «говорить»: что он означает для университета, для студентов?",
        "— Используй академические конструкции живо, не механически:",
        "  «это не просто статистика — за ней стоит...», «показательно, что...»,",
        "  «такое сочетание сигнализирует о...», «вместе с тем картина не однородна».",
        "— Формулируй выводы конкретно — не «необходимо улучшить», а что именно и почему.",
        "— Все слова не на русском языке переводи на русский.",
        "— Объём: 6–10 предложений.",
        "",
    ]

    if sec_name:
        lines += [f"Раздел анкеты: «{sec_name}»", ""]

    q = question
    lines += [
        f"Напиши аналитический фрагмент по вопросу: «{q['question_name']}»",
        "",
        "Статистика ответов:",
    ]
    for row in q["rows"]:
        parts = []
        for fk in q["file_keys"]:
            label = q["file_labels"].get(fk, fk)
            count = row["counts"].get(fk, 0)
            total = q["file_totals"].get(fk, 0)
            pct = round(count / total * 100, 1) if total > 0 else 0
            parts.append(f"{label}: {count} ({pct}%)")
        lines.append(f"  - {row['answer']}: {'; '.join(parts)}")

    if sec_description:
        lines += [
            "",
            "При написании учти контекст раздела, указанный в начале.",
        ]

    user_prompt = "\n".join(lines)
    return system_prompt, user_prompt


# ===================== ВЫЗОВ МОДЕЛИ =====================

def _call_openrouter(prompt: str, api_key: str, system: str = "") -> str:
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    for attempt in range(5):
        try:
            response = client.chat.completions.create(
                model=OPENROUTER_MODEL,
                messages=messages,
                max_tokens=3000,
                temperature=0.7,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if "429" in str(e) and attempt < 4:
                wait = 15 * (attempt + 1)
                print(f"  -> Rate limit, жду {wait}с...")
                time.sleep(wait)
            else:
                raise


def _call_llm(prompt: str, system: str = "") -> str:
    key = OPENROUTER_API_KEY_DEFAULT
    if not key:
        raise ValueError("OPENROUTER_API_KEY не задан в .env")
    return _call_openrouter(prompt, key, system=system)


# ===================== АНАЛИТИКА (файл 2) =====================

def generate_analysis_docx(questions: list, progress_callback=None) -> bytes:
    """Аналитический файл: один вызов LLM на каждый вопрос."""
    doc = _make_doc()
    total_questions = len(questions)
    _last_section = None
    chart_counter = [1]

    for idx, q in enumerate(questions, start=1):
        sec = q.get("section")
        sec_name = sec.get("name") if sec else ""
        sec_description = sec.get("description", "") if sec else ""

        if progress_callback:
            progress_callback(idx, total_questions, q["question_name"])

        print(f"[{idx}/{total_questions}] Генерация: {q['question_name']}")

        if sec_name and sec_name != _last_section:
            _last_section = sec_name
            if idx > 1:
                doc.add_page_break()
            _p(doc, sec_name, bold=True, size=14, space_after=4)
            if sec_description:
                _p(doc, sec_description, size=11, space_after=6)

        _p(
            doc,
            f"Вопрос {q['table_num']} — «{q['question_name']}»",
            bold=True,
            size=12,
            space_before=8,
            space_after=3,
        )

        try:
            system_prompt, user_prompt = _build_question_prompt(q, sec_name, sec_description)
            analysis = _call_llm(user_prompt, system=system_prompt)
            _p(doc, analysis, space_after=6)
            print("  -> OK")
            time.sleep(10)
        except Exception as e:
            print(f"  -> ERROR: {e}")
            _p(doc, f"Ошибка генерации аналитики: {e}", bold=True, space_after=4)

        insert_visualization(doc, q, chart_counter)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ===================== ЕДИНАЯ ТОЧКА ВХОДА =====================

def generate_docx(questions: list, progress_callback=None) -> tuple[bytes, bytes]:
    data_bytes = generate_data_docx(questions)
    analysis_bytes = generate_analysis_docx(questions, progress_callback=progress_callback)
    return data_bytes, analysis_bytes
