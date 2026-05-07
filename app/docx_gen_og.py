import io
from docx import Document
from docx.shared import Pt, Cm

#ФУНКЦИЯ ЭКСПОРТА (Редактировать тут)

def _p(doc, text, bold=False, size=12, space_before=0, space_after=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return para


def generate_docx(questions: list) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for q in questions:
        file_keys = q["file_keys"]
        file_labels = q["file_labels"]
        file_totals = q["file_totals"]
        is_single = len(file_keys) == 1

        # Question heading
        _p(doc,
           f"Вопрос {q['table_num']} – «{q['question_name']}»",
           bold=True, space_before=10, space_after=2)

        # One line per answer
        for row in q["rows"]:
            if is_single:
                fk = file_keys[0]
                count = row["counts"].get(fk, 0)
                total = file_totals.get(fk, 0)
                pct = f"{count / total * 100:.1f}%" if total > 0 else "—"
                _p(doc, f"  • {row['answer']}: {count} ({pct})", space_after=1)
            else:
                parts = []
                for fk in file_keys:
                    label = file_labels.get(fk, fk)
                    count = row["counts"].get(fk, 0)
                    total = file_totals.get(fk, 0)
                    pct = f"{count / total * 100:.1f}%" if total > 0 else "—"
                    parts.append(f"{label}: {count} ({pct})")
                _p(doc, f"  • {row['answer']}: {'; '.join(parts)}", space_after=1)

        # Total line
        if q.get("show_total", True):
            if is_single:
                fk = file_keys[0]
                _p(doc, f"  Всего: {file_totals.get(fk, 0)}", bold=True, space_after=6)
            else:
                parts = [f"{file_labels.get(fk, fk)}: {file_totals.get(fk, 0)}" for fk in file_keys]
                _p(doc, f"  Всего: {'; '.join(parts)}", bold=True, space_after=6)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()