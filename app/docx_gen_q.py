import io
import ollama

from docx import Document
from docx.shared import Pt, Cm


# =========================
# WORD HELPERS
# =========================

def _p(doc, text, bold=False, size=12, space_before=0, space_after=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)

    return para


# =========================
# AI ANALYSIS
# =========================

def build_prompt(q):
    """
    Формирует промпт для нейросети
    """

    lines = []

    lines.append(
        f"Вопрос анкеты: {q['question_name']}\n"
    )

    lines.append("Статистика ответов:")

    for row in q["rows"]:
        answer = row["answer"]

        parts = []

        for fk in q["file_keys"]:
            label = q["file_labels"].get(fk, fk)

            count = row["counts"].get(fk, 0)
            total = q["file_totals"].get(fk, 0)

            pct = round(count / total * 100, 1) if total > 0 else 0

            parts.append(
                f"{label}: {count} ({pct}%)"
            )

        lines.append(
            f"- {answer}: " + "; ".join(parts)
        )

    lines.append("""
Сделай аналитический вывод по результатам вопроса. Пиши официальным стилем.
Вывод должен быть одним абзацем.
Если встречаются слова не на русском языке, переводи их на русский.                 
Сделай вывод о тенденциях и распределении ответов.
Выделяй уникальные значения, особенно если это вопрос с открытым ответом.
Объем: 4-6 предложений.
""")

    return "\n".join(lines)


def generate_analysis(q):
    """
    Генерация аналитики через Ollama
    """

    prompt = build_prompt(q)

    response = ollama.chat(
        model="mistral",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response["message"]["content"].strip()


# =========================
# DOCX EXPORT
# =========================

def generate_docx(questions: list) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    total_questions = len(questions)

    for index, q in enumerate(questions, start=1):

        print(
            f"[{index}/{total_questions}] "
            f"Генерация аналитики: "
            f"Вопрос {q['table_num']} "
            f"— {q['question_name']}"
        )

        file_keys = q["file_keys"]
        file_labels = q["file_labels"]
        file_totals = q["file_totals"]

        is_single = len(file_keys) == 1

        # =========================
        # Заголовок
        # =========================

        _p(
            doc,
            f"Вопрос {q['table_num']} – «{q['question_name']}»",
            bold=True,
            space_before=10,
            space_after=2
        )

        # =========================
        # Таблица ответов
        # =========================

        for row in q["rows"]:

            if is_single:

                fk = file_keys[0]

                count = row["counts"].get(fk, 0)
                total = file_totals.get(fk, 0)

                pct = (
                    f"{count / total * 100:.1f}%"
                    if total > 0 else "—"
                )

                _p(
                    doc,
                    f"  • {row['answer']}: {count} ({pct})",
                    space_after=1
                )

            else:

                parts = []

                for fk in file_keys:

                    label = file_labels.get(fk, fk)

                    count = row["counts"].get(fk, 0)
                    total = file_totals.get(fk, 0)

                    pct = (
                        f"{count / total * 100:.1f}%"
                        if total > 0 else "—"
                    )

                    parts.append(
                        f"{label}: {count} ({pct})"
                    )

                _p(
                    doc,
                    f"  • {row['answer']}: {'; '.join(parts)}",
                    space_after=1
                )

        # =========================
        # Итого
        # =========================

        if q.get("show_total", True):

            if is_single:

                fk = file_keys[0]

                _p(
                    doc,
                    f"  Всего: {file_totals.get(fk, 0)}",
                    bold=True,
                    space_after=4
                )

            else:

                parts = [
                    f"{file_labels.get(fk, fk)}: {file_totals.get(fk, 0)}"
                    for fk in file_keys
                ]

                _p(
                    doc,
                    f"  Всего: {'; '.join(parts)}",
                    bold=True,
                    space_after=4
                )

        # =========================
        # AI аналитика
        # =========================

        try:

            analysis = generate_analysis(q)

            _p(
                doc,
                "Аналитический вывод:",
                bold=True,
                space_before=2,
                space_after=1
            )

            _p(
                doc,
                analysis,
                space_after=8
            )

            print("  -> OK")

        except Exception as e:

            print(f"  -> ERROR: {e}")

            _p(
                doc,
                "Ошибка генерации аналитики.",
                bold=True,
                space_after=8
            )

    buf = io.BytesIO()

    doc.save(buf)

    buf.seek(0)

    return buf.getvalue()